#!/usr/bin/env python3
"""
Technical Report Generator (DOCX)
=================================
Generates an elaborate .docx report with properly formatted equations.
"""
import os
import sys
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Application imports
sys.path.append(os.getcwd())
from quantum_vascular_coils import QUANTUM_VASCULAR_COIL_LIBRARY
from statistical_adaptive_pulse import ADAPTIVE_SEQUENCES

def add_equation(paragraph, latex_content):
    """
    Simulates equation formatting. 
    True LaTeX to OMML is complex without external libs.
    We'll format it as a centered, distinct block.
    """
    run = paragraph.add_run(latex_content)
    run.font.name = 'Cambria Math'
    run.font.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 50, 100)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def parse_markdown_to_docx(doc, md_content):
    lines = md_content.split('\n')
    in_code_block = False
    in_math_block = False
    math_buffer = []

    for line in lines:
        line = line.strip()
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            
        # Math Blocks
        elif line.startswith('$$') and line.endswith('$$'):
            p = doc.add_paragraph()
            add_equation(p, line[2:-2].strip())
        elif line.startswith('$$'):
            in_math_block = True
            math_buffer = [line[2:]]
        elif line.endswith('$$') and in_math_block:
            in_math_block = False
            math_buffer.append(line[:-2])
            p = doc.add_paragraph()
            add_equation(p, '\n'.join(math_buffer))
            math_buffer = []
        elif in_math_block:
            math_buffer.append(line)
            
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
            
        # Standard
        elif line == '':
            continue
        elif line.startswith('---'):
            doc.add_page_break()
        else:
            # Basic formatting removal for clean text
            clean_line = line.replace('**', '').replace('*', '').replace('`', '')
            doc.add_paragraph(clean_line)

def generate_docx():
    doc = Document()
    
    # Title
    title = doc.add_heading('Neurovascular Coil Technical Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph('Finite Math Derivations & Pulse Sequences')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_page_break()
    
    # Part 1: Theory
    md_path = "Quantum_Vascular_Coil_Theory.md"
    if os.path.exists(md_path):
        with open(md_path, 'r') as f:
            parse_markdown_to_docx(doc, f.read())
            
    doc.add_page_break()
    
    # Part 1.5: Supplement
    supp_path = "Quantum_Advanced_Pulse_Sequences.md"
    if os.path.exists(supp_path):
        doc.add_heading("Supplement: Advanced Quantum Sequences", level=1)
        with open(supp_path, 'r') as f:
            parse_markdown_to_docx(doc, f.read())
            
    doc.add_page_break()
    
    # Part 2: Coils
    doc.add_heading("Part II: Coil Implementation & Derivations", level=1)
    
    for idx, coil_class in sorted(QUANTUM_VASCULAR_COIL_LIBRARY.items()):
        coil = coil_class()
        doc.add_heading(f"{idx}. {coil.name}", level=2)
        
        # Specs
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parameter'
        hdr_cells[1].text = 'Value'
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'Name'
        row_cells[1].text = str(coil.name)
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'Elements'
        row_cells[1].text = str(coil.num_elements)
        
        # Derivation
        doc.add_heading("Mathematical Derivation:", level=3)
        if coil_class.__doc__:
            math_text = coil_class.__doc__.strip()
            p = doc.add_paragraph()
            add_equation(p, math_text)
            
        doc.add_paragraph() # Spacer
        
    doc.add_page_break()
    
    # Part 3: Pulse Sequences
    doc.add_heading("Part III: Pulse Sequence Math", level=1)
    
    for seq_id, seq_class in ADAPTIVE_SEQUENCES.items():
        seq = seq_class()
        doc.add_heading(f"Sequence: {seq.sequence_name}", level=2)
        
        if seq_class.__doc__:
            doc.add_heading("Description & Math:", level=3)
            p = doc.add_paragraph()
            add_equation(p, seq_class.__doc__.strip())
            
    filename = "Neurovascular_Coil_Technical_Report.docx"
    doc.save(filename)
    print(f"DOCX Report generated: {filename}")

if __name__ == "__main__":
    generate_docx()
