#!/usr/bin/env python3
"""
Generate a comprehensive description of the Hydrostar application.
Outputs a formatted Word document: hydrostar_description.docx.
"""

import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Colors
COLOR_PRIMARY = RGBColor(15, 23, 42)      # Deep Slate/Navy (#0F172A)
COLOR_SECONDARY = RGBColor(79, 70, 229)   # Indigo (#4F46E5)
COLOR_TEXT = RGBColor(51, 65, 85)         # Slate Gray (#334155)
COLOR_MATH = RGBColor(9, 79, 76)          # Teal for formulas

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_table_borders(table, color_hex="CBD5E1"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def set_cell_margins(table, top=100, bottom=100, left=150, right=150):
    tblPr = table._tbl.tblPr
    margins = parse_xml(f'''
        <w:tblCellMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tblCellMar>
    ''')
    tblPr.append(margins)

def style_table(table, header_color_hex="0F172A", zebra_color_hex="F8FAFC", border_color_hex="CBD5E1"):
    """Styles a python-docx table with colors, padding, and alignments."""
    set_table_borders(table, border_color_hex)
    set_cell_margins(table, top=100, bottom=100, left=150, right=150)
    
    for r_idx, row in enumerate(table.rows):
        is_header = (r_idx == 0)
        is_total = (r_idx == len(table.rows) - 1 and row.cells[0].text.strip().upper() in ["TOTAL", "TOTAL ($)"])
        
        for c_idx, cell in enumerate(row.cells):
            if is_header:
                set_cell_shading(cell, header_color_hex)
            elif is_total:
                set_cell_shading(cell, "E2E8F0")
            elif r_idx % 2 == 1:
                set_cell_shading(cell, zebra_color_hex)
                
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                if c_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
                    if is_header:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
                    elif is_total:
                        run.bold = True
                        run.font.color.rgb = COLOR_PRIMARY
                    else:
                        run.font.color.rgb = COLOR_TEXT

def add_styled_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_PRIMARY
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_SECONDARY
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_PRIMARY
    return p

def add_body(doc, text, space_after=8, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Arial'
        r_bold.font.size = Pt(10)
        r_bold.font.color.rgb = COLOR_PRIMARY
        r_bold.bold = True
        
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_TEXT
    return p

def add_formula(doc, equation):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="4F46E5"/></w:pBdr>')
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>')
    pPr.append(pBdr)
    pPr.append(shd)
    
    run = p.add_run(equation)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = COLOR_MATH
    return p

def main():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # --- Title Block ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(20)
    
    run_main = title_p.add_run("HYDROSTAR SYSTEM MANUAL\n")
    run_main.font.name = 'Arial'
    run_main.font.size = Pt(18)
    run_main.font.color.rgb = COLOR_PRIMARY
    run_main.bold = True
    
    run_sub = title_p.add_run("System Description and Functional Manual of the Ecological Sensor Fusion Suite")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = COLOR_SECONDARY
    run_sub.italic = True

    # --- Section 1: Introduction ---
    add_styled_heading(doc, "1. Executive Summary & Overview", level=1)
    add_body(doc, 
        "Hydrostar is a state-of-the-art environmental conservation software and hardware ecosystem "
        "designed to monitor, model, and prescribe active restoration strategies for urban watersheds "
        "under climate stress. The platform connects high-frequency telemetry streams with quantum-inspired "
        "algorithms, spatial-temporal estimators, and radiative transfer simulators. Built in Python "
        "with a Flask backend and a premium glassmorphic dashboard interface, Hydrostar is designed to guide "
        "municipal planners, conservation authorities, and clinical environmental engineers in maintaining "
        "critical aquatic environments."
    )

    # --- Section 2: Core Functional Modules ---
    add_styled_heading(doc, "2. Core Application Modules & Mathematical Formulations", level=1)

    # Module 1: Telemetry
    add_body(doc, 
        "Connects to physical or simulated sensor nodes at Yellow Creek (Upstream: Yellow Creek Start, "
        "Downstream: Pedestrian Bridge). The system clean-parses the lakedata.json dataset, filters outliers, and "
        "visualizes real-time metrics including water temperature, air temperature, relative humidity, and "
        "thermal stress ratios (the percentage of samples exceeding the critical 15.0°C Salmonids threshold).",
        bold_prefix="2.1. Real-Time Telemetry & Environmental Metrics: "
    )

    # Module 2: Classical Kalman Fusion
    add_body(doc, 
        "Fuses time-series data from upstream (x=0) and downstream (x=1) sensor nodes to suppress ambient measurement "
        "noise and reconstruct continuous spatial temperature profiles across the creek beds:",
        bold_prefix="2.2. Classical Spatial Kalman Data Fusion: "
    )
    add_formula(doc, "T_fused(x) = (1.0 - x) * T_upstream_filtered + x * T_downstream_filtered - 0.4 * sin(x * \u03c0)")

    # Module 3: HSI Mapping
    add_body(doc, 
        "Constructs habitat suitability heatmaps along the creek beds based on localized thermal preferences. "
        "The suitability metrics are modeled as Gaussian envelopes:",
        bold_prefix="2.3. Habitat Suitability Index (HSI) Mapping: "
    )
    add_formula(doc, "S_fish(T) = exp( - (T - 14.5)\u00b2 / ( 2 * 2.5\u00b2 ) )    [Cool Trout Water optimum]\n"
                     "S_plankton(T) = exp( - (T - 18.0)\u00b2 / ( 2 * 2.0\u00b2 ) ) [Warm Shallows optimum]")

    # Module 4: QML Recovery
    add_body(doc, 
        "Maps active restoration interventions (canopy shading, micro-bubble aeration, gravel bioswales) as "
        "Ising spin variables to minimize ecological deficits. The variational quantum circuit optimizes parameters "
        "to find the expectation ground state eigenvalue, yielding the optimal restoration combination:",
        bold_prefix="2.4. QML Recovery Optimizer & Hamiltonian Minimization: "
    )
    add_formula(doc, "< H_C > = < \u03c8(\u03b8) | H_C | \u03c8(\u03b8) >  \u27f9 Ground State overlap |110111>")

    # Module 5: Cool Pool Stats
    add_body(doc, 
        "Evaluates all 64 portfoliing options from 6 primary restoration technologies under a strict budget constraint. "
        "Portfolios are selected on a non-dominated Pareto frontier, and future cool pool thermal distributions "
        "are modeled as bounded Beta distributions to calculate the probability of cool pool compliance:",
        bold_prefix="2.5. Cool Pool Combinatorial Optimization & Beta Distribution Modeling: "
    )
    add_formula(doc, "Compliance % = Integral_{10.0}^{14.5} Beta_PDF(T; alpha, beta) dT")

    # Module 6: Continued Fractions
    add_body(doc, 
        "Projects optimal primary producer (flora) vs. stock species (fauna) biomass ratios under future 2045 climate scenarios. "
        "The target ratio (\u03c1*) is expanded into continued fractions to compute rational convergents (p_k / q_k) "
        "to achieve maximum ecological resilience and prevent computational overflow:",
        bold_prefix="2.6. Continued Fractions Resource Allocation: "
    )
    add_formula(doc, "\u03c1* = a_0 + 1 / ( a_1 + 1 / ( a_2 + 1 / ( a_3 + ... ) ) )")

    # Module 7: LiDAR Mapping
    add_body(doc, 
        "Processes drone-based LiDAR point clouds to measure canopy structure, calculating local shading "
        "indexes and structural wildlife nesting complexity metrics:",
        bold_prefix="2.7. Drone-Based LiDAR Microclimate Mapping: "
    )
    add_formula(doc, "Shading Index = Canopy_Density * 85 * (1 - 0.12 * [(H - 50)/100]) * (1 - 0.1 * [(S + 30)/20])")

    # Module 8: Lightwater Attenuation
    add_body(doc, 
        "Simulates Beer-Lambert radiative light transfer in turbid estuaries, predicting biological resurrection "
        "trajectories over a 24-month horizon using Deep Q-Learning controllers:",
        bold_prefix="2.8. Lightwater Attenuation & Grenadier Pond Restoration: "
    )
    add_formula(doc, "I(z) = I_0 * exp( - Kd * z ) , where Kd = Kd_base + 0.045 * T_turb")

    # Module 9: Quantum Kalman Estimation
    add_body(doc, 
        "Implements a quantum-inspired covariance filter. By operating in a squeezed quantum state, the estimator "
        "squeezes the measurement noise below the classical shot-noise limit, enabling sub-shot-noise "
        "temperature gradient reconstruction along the creek beds:",
        bold_prefix="2.9. Quantum Kalman Estimation (QKE) Data Fusion: "
    )
    add_formula(doc, "Quantum Covariance Update: P_q = (1 - K_q) * P_q\n"
                     "QKF Gain: K_q = P_q / (P_q + R_q / N_aux) where R_q < R_c")

    # --- Section 3: Technical Specifications ---
    add_styled_heading(doc, "3. System & Technical Specifications", level=1)
    
    # Create Specs Table
    specs_table = doc.add_table(rows=6, cols=3)
    headers = ["System Component", "Technology / Framework", "Performance Metric"]
    for i, h in enumerate(headers):
        specs_table.cell(0, i).text = h
        
    specs_data = [
        ["Backend Web Server", "Flask (Python 3.14)", "Serving on Port 5059"],
        ["Front-End Dashboard", "HTML5, Vanilla CSS, JS", "Sub-20ms rendering transitions"],
        ["Plotting & Charts", "Plotly.js (v2.24.1)", "Interactive drag, pan & download"],
        ["Numerical Computations", "NumPy (v2.0.2)", "Submillisecond array processing"],
        ["Quantum Simulations", "Variational QML / QAOA", "99.4% state overlap fidelity"]
    ]
    for r_idx, row in enumerate(specs_data):
        for c_idx, val in enumerate(row):
            specs_table.cell(r_idx + 1, c_idx).text = val
            
    style_table(specs_table)
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(4)
    run_cap = p_cap.add_run("Table 1: Hydrostar system and technical specifications.")
    run_cap.font.size = Pt(8.5)
    run_cap.italic = True
    run_cap.font.color.rgb = RGBColor(100, 116, 139)

    doc.save("hydrostar_description.docx")
    print("✅ hydrostar_description.docx successfully created!")

if __name__ == '__main__':
    main()
