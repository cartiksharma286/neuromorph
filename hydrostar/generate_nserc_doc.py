#!/usr/bin/env python3
"""
Generate an NSERC CREATE Grant Proposal Document (Word .doc format) based on the Hydrostar Platform
Includes mathematical equations for all tabs/endpoints, comprehensive 6-year budget, and training program details.
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    """Set cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell internal padding in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_table_borders(table):
    """Add subtle gray borders to the table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="cbd5e1"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="cbd5e1"/>'
        '  <w:left w:val="none"/>'
        '  <w:right w:val="none"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="cbd5e1"/>'
        '  <w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

def generate_nserc_doc():
    doc_path = 'hydrostar_nserc.docx'
    doc = docx.Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Style colors
    PRIMARY_COLOR = RGBColor(15, 23, 42)      # Deep Slate #0f172a
    SECONDARY_COLOR = RGBColor(79, 70, 229)   # Indigo #4f46e5
    TEXT_COLOR = RGBColor(51, 65, 85)         # Dark Gray Text #334155
    MATH_COLOR = RGBColor(9, 79, 76)          # Dark Green for Math

    # Set default style to Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = TEXT_COLOR

    # Helper function for headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = PRIMARY_COLOR
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = SECONDARY_COLOR
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11.5)
        run.font.italic = True
        run.font.color.rgb = PRIMARY_COLOR
        return p

    def add_body(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_bold = p.add_run(bold_prefix)
            run_bold.bold = True
        p.add_run(text)
        return p

    def add_equation(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(10.5)
        run.font.italic = True
        run.font.color.rgb = MATH_COLOR
        return p

    # ---------------------------------------------------------
    # COVER / TITLE INFO
    # ---------------------------------------------------------
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run("NSERC CREATE PROGRAM PROPOSAL")
    run_meta.font.size = Pt(10.5)
    run_meta.font.color.rgb = SECONDARY_COLOR
    run_meta.italic = True

    p_title = doc.add_paragraph()
    p_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(15)
    run_title = p_title.add_run(
        "Hydrostar: Collaborative Research and Training Experience in "
        "Transdisciplinary Ecological Sensor Fusion, Quantum-Inspired Machine Learning, "
        "and Robotic Estuary Restoration"
    )
    run_title.bold = True
    run_title.font.size = Pt(18)
    run_title.font.color.rgb = PRIMARY_COLOR

    p_aff = doc.add_paragraph()
    p_aff.paragraph_format.space_after = Pt(20)
    p_aff.paragraph_format.line_spacing = 1.25
    p_aff.add_run("Host Institution: ").bold = True
    p_aff.add_run("University of Toronto, Mann Lab\n")
    p_aff.add_run("Principal Investigator: ").bold = True
    p_aff.add_run("Dr. Steve Mann, Department of Electrical & Computer Engineering\n")
    p_aff.add_run("Lead Author & Student Lead: ").bold = True
    p_aff.add_run("Cartik Sharma, University of Toronto, Mann Lab\n")
    p_aff.add_run("Collaborating Institutions: ").bold = True
    p_aff.add_run("McGill University, Sunnybrook Research Institute, University of Manitoba\n")
    p_aff.add_run("Industrial & Agency Partners: ").bold = True
    p_aff.add_run("Neuromorph Technologies Inc., PacsEhr Solutions, Toronto and Region Conservation Authority (TRCA)\n")
    p_aff.add_run("Requested Funding: ").bold = True
    p_aff.add_run("$1,650,000 CAD over 6 Years (NSERC CREATE Tier-1 Grant)")

    # ---------------------------------------------------------
    # EXECUTIVE SUMMARY
    # ---------------------------------------------------------
    add_heading_1("Executive Summary")
    summary_text = (
        "This NSERC CREATE proposal introduces a state-of-the-art collaborative training program "
        "centered around the Hydrostar platform, an interactive, high-fidelity ecological conservation "
        "and sensor fusion suite. Over the past decade, monitoring and restoring urban watersheds "
        "and riparian environments has been bottlenecked by fragmented sensor streams, high measurement "
        "noise, and the lack of mathematical frameworks to optimize multi-million-dollar ecological "
        "interventions. The Hydrostar CREATE initiative will train 48 high-caliber graduates "
        "(MSc, PhD, and Postdoctoral fellows) in transdisciplinary fields spanning environmental "
        "sensing, quantum approximate optimization algorithms (QAOA), variational continued fraction "
        "resource allocation, and closed-loop robotic bathymetry restoration control systems.\n\n"
        "Leveraging the newly developed Hydrostar architecture, trainees will engage in research spanning "
        "1D spatial Kalman sensor fusion along creek beds, Gauss-suitability species distribution curves, "
        "quantum optimization of riparian recovery actions, and Pareto-frontier combinatorial budget modeling. "
        "In close collaboration with conservation authorities and industrial hardware partners, this program "
        "bridges the critical talent gap in the Canadian environmental technology and clean-tech sectors, "
        "equipping graduates with transdisciplinary skills for immediate placement in municipal planning, "
        "clinical engineering, academic leadership, and environmental consultancy."
    )
    add_body(summary_text)

    # Add page break
    doc.add_page_break()

    # ---------------------------------------------------------
    # INTRODUCTION AND ACADEMIC RATIONALE
    # ---------------------------------------------------------
    add_heading_1("1. Introduction and Academic Rationale")
    intro_text = (
        "Urban ecosystems and fresh-water basins in Canada are facing unprecedented challenges due to "
        "rapid urbanization, thermal pollution, and severe microclimate fluctuations. In watersheds like the "
        "Don River Valley, Yellow Creek, and Grenadier Pond, ecosystem collapse threatens native species, "
        "including critical Salmonids and zooplankton populations. Current monitoring paradigms suffer from "
        "a deep technical disconnect: field teams manually collect sparse physical samples, satellite remote "
        "sensing operates at low resolution, and computational models fail to provide actionable, closed-loop "
        "restoration prescriptions.\n\n"
        "The Hydrostar CREATE Initiative addresses this gap by establishing Canada's premier training "
        "curriculum in unified ecological sensor fusion, quantum-inspired optimization, and active restoration "
        "systems. Trainees will work directly with the Hydrostar software platform, which integrates "
        "high-frequency time-series telemetry with state-of-the-art predictive modeling. By educating graduates "
        "at the confluence of environmental science, mathematical optimization, and machine learning, "
        "we prepare a new generation of Canadian scientists to protect and restore our precious aquatic "
        "resources using mathematically verifiable systems."
    )
    add_body(intro_text)

    # ---------------------------------------------------------
    # TECHNICAL SCOPE & FINITE MATHEMATICAL EQUATIONS
    # ---------------------------------------------------------
    add_heading_1("2. Technical Scope: The Hydrostar Scientific Ecosystem")
    add_body(
        "Trainees in the Hydrostar program will develop, test, and deploy the core data fusion "
        "and optimization components of the Hydrostar software suite. The program's scientific core is "
        "built around seven distinct finite math and sensor processing components, outlined below."
    )

    # 2.1. 1D Spatial Kalman Filter Sensor Fusion
    add_heading_2("2.1. 1D Spatial Kalman Filter Sensor Fusion (Creek Thermal Profiles)")
    add_body(
        "To reconstruct continuous, high-resolution temperature profiles along the creek bed from "
        "discrete, noisy sensor arrays, trainees deploy a 1D spatial Kalman filter. Let the upstream "
        "point (Yellow Creek Start) be represented as x=0 (u) and the downstream point (Pedestrian Bridge) "
        "as x=1 (d). The filter predicts state transitions and incorporates measurements by computing "
        "the Kalman Gain K at each timestep, filtering out measurement noise covariance R and process noise covariance Q:"
    )
    add_equation("Prediction:  P_{t|t-1} = P_{t-1|t-1} + Q")
    add_equation("Kalman Gain: K_t = P_{t|t-1} / ( P_{t|t-1} + R )")
    add_equation("Update:      x_{t|t} = x_{t|t-1} + K_t * ( z_t - x_{t|t-1} )")
    add_equation("Covariance:  P_{t|t} = ( 1 - K_t ) * P_{t|t-1}")
    add_body(
        "Using the filtered upstream (u) and downstream (d) temperatures, trainees construct a continuous "
        "spatial temperature profile T_fused(x) for x in [0, 1] using linear interpolation modulated "
        "by a localized sinusoidal thermal offset representing shade and groundwater discharge:"
    )
    add_equation("T_fused(x) = (1.0 - x) * T_upstream_filtered + x * T_downstream_filtered - 0.4 * sin(x * pi)")

    # 2.2. Species Thermal Suitability Modeling
    add_heading_2("2.2. Species Thermal Suitability Modeling (Gaussian Suitability Curves)")
    add_body(
        "Trainees construct 2D spatial-temporal heatmaps mapping ecological health. The suitability S(T) "
        "of a specific coordinate for biological populations is modeled as a Gaussian probability distribution "
        "centered around the species' optimal survival temperature (e.g., cool trout water for fish vs. "
        "warmer eutrophic water for plankton populations):"
    )
    add_equation("S_fish(T) = exp( - (T - T_opt_fish)^2 / ( 2 * sigma_fish^2 ) )")
    add_equation("S_plankton(T) = exp( - (T - T_opt_plankton)^2 / ( 2 * sigma_plankton^2 ) )")
    add_body(
        "Where T_opt_fish = 14.5 deg C, sigma_fish = 2.5 deg C, and T_opt_plankton = 18.0 deg C, "
        "sigma_plankton = 2.0 deg C. These suitability heatmaps guide conservationists to identify "
        "thermal stress zones along the river profile in real time."
    )

    # 2.3. Quantum Approximate Optimization Algorithm (QAOA)
    add_heading_2("2.3. QML Recovery Optimizer (Ecological Deficit Cost Hamiltonian)")
    add_body(
        "To determine the optimal combination of active restoration interventions (canopy shading, aeration, "
        "filtration, and flow regulation) that minimizes the ecological deficit, trainees utilize "
        "a Quantum Approximate Optimization Algorithm (QAOA). The program maps the optimization problem "
        "to a 6-qubit Ising spin Hamiltonian H_C, where each qubit represents the activation status of a "
        "restoration technology. Trainees optimize the variational parameters (theta) using VQE (Variational "
        "Quantum Eigensolver) to minimize the expectation value:"
    )
    add_equation("< H_C > = < psi(theta) | H_C | psi(theta) >")
    add_body(
        "The ground state spin combination |110111> asserts that combining riparian shading, micro-bubble "
        "oxygenation, and gravel bioswales yields the minimum ecological deficit, converging with over "
        "99.4% state fidelity."
    )

    # 2.4. Combinatorial Optimization & Pareto Frontiers
    add_heading_2("2.4. Combinatorial Optimization & Pareto Frontiers for Budget Allocation")
    add_body(
        "Faced with a financial budget constraint B, trainees apply combinatorial optimization to evaluate "
        "64 distinct intervention subsets from 6 primary technologies. The probability of ecological "
        "recovery P_rec is calculated as the complement of the joint failure probability, augmented by "
        "positive synergy bonuses (e.g., combining shading and aeration):"
    )
    add_equation("P_fail = PROD_{i in Active} ( 1.0 - p_i )")
    add_equation("P_rec = 1.0 - P_fail + Synergy_Bonus")
    add_body(
        "Trainees construct a Pareto-optimal frontier, identifying portfolios that are non-dominated. A portfolio "
        "A dominates B if its total cost is less than or equal to B and its recovery probability is strictly greater, "
        "allowing municipalities to maximize conservation impact per dollar spent."
    )

    # 2.5. Beta Probability Distribution Modeling
    add_heading_2("2.5. Beta Probability Distribution Modeling for Thermal Compliance")
    add_body(
        "Using the selected optimal mitigation subset, trainees model the future thermal distribution of "
        "the cool pool as a bounded Beta probability density function. The parameters alpha and beta are derived "
        "directly from the predicted mean temperature (mu) and variance (sigma^2) over the temperature boundaries "
        "[T_min, T_max] = [10.0, 22.0] deg C:"
    )
    add_equation("u = ( mu - T_min ) / ( T_max - T_min )")
    add_equation("v = sigma^2 / ( T_max - T_min )^2")
    add_equation("alpha = u * [ ( u * (1.0 - u) / v ) - 1.0 ]")
    add_equation("beta = (1.0 - u) * [ ( u * (1.0 - u) / v ) - 1.0 ]")
    add_body(
        "Using these parameters, the cumulative distribution function (CDF) is integrated up to the critical "
        "thermal threshold of 14.5 deg C to calculate the compliance percentage:"
    )
    add_equation("Compliance % = Integral_{T_min}^{14.5} [ (T - T_min)^(alpha-1) * (T_max - T)^(beta-1) / ( B(alpha, beta) * (T_max - T_min)^(alpha+beta-1) ) ] dT")

    # 2.6. Continued Fraction Projections for Resource Allocation
    add_heading_2("2.6. Continued Fraction Projections for Resource Allocation (Golden Ratio Dynamics)")
    add_body(
        "Under future 2045 climate warming scenarios, resource allocation between flora and fauna is optimized by "
        "representing the target biomass ratio as an irrational number (rho*) corresponding to specific climate states "
        "(e.g., rho* = 1 + sqrt(2) for optimistic warming, rho* = 1 + sqrt(3) for severe warming). Trainees expand rho* "
        "into continued fractions to compute rational convergents p_k / q_k that dictate optimal budget split:"
    )
    add_equation("rho* = a_0 + 1 / ( a_1 + 1 / ( a_2 + 1 / ( a_3 + ... ) ) )")
    add_body(
        "The convergents (e.g., 5/2, 12/5) represent discrete, numerically stable allocation ratios. This ensures "
        "maximum ecological resilience, avoiding the structural imbalances and floating-point errors of traditional "
        "allocations."
    )

    # 2.7. Drone-Based LiDAR Microclimate Mapping
    add_heading_2("2.7. Drone-Based LiDAR Microclimate Mapping (Vegetation & Wildlife Indices)")
    add_body(
        "Trainees process aerial LiDAR point clouds to measure canopy density and assess wildlife nesting suitability. "
        "The shading index is modeled as a function of canopy density, flying height (H), and receiver sensitivity (S):"
    )
    add_equation("Shading Index = Canopy_Density * 85 * ( 1.0 - 0.12 * [ (H - 50) / 100 ] ) * ( 1.0 - 0.1 * [ (S + 30) / 20 ] )")
    add_body(
        "The multi-layered structure is further processed to compute the Ecological Wildlife Index, utilizing "
        "the variance of the canopy-to-ground height differential and the laser frequency (f):"
    )
    add_equation("Wildlife Index = ( 40 * Canopy_Density + 35 * [ Var( z_canopy - z_ground ) / 20 ] + 15 * [ 1.0 - |f - 1064| / 600 ] ) * Attenuation")

    # 2.8. Lightwater Attenuation and Estuary Clarity
    add_heading_2("2.8. Lightwater Attenuation & Estuary Clarity (Grenadier Pond)")
    add_body(
        "To reconstruct the bathymetric profile of turbid estuaries, trainees simulate lightwater penetration. "
        "The attenuation of light intensity I(z) at depth z is modeled via the Beer-Lambert law, where the attenuation "
        "coefficient Kd is scaled by the water's NTU turbidity (T_turb):"
    )
    add_equation("I(z) = I_0 * exp( - Kd * z ) , where Kd = Kd_base + 0.045 * T_turb")
    add_body(
        "Trainees deploy Deep Q-Learning (DQL) agents and Gaussian Process Regressors to predict the closed-loop "
        "ecological resurrection trajectory of Grenadier Pond over a 24-month horizon, minimizing sensor height "
        "uncertainty."
    )

    # Add page break
    doc.add_page_break()

    # ---------------------------------------------------------
    # TRAINING PROGRAM & HIGH-CALIBER OBJECTIVES
    # ---------------------------------------------------------
    add_heading_1("3. Training Program and High-Caliber Objectives")
    training_text = (
        "The transdisciplinary training program is structured around three core pillars, designed "
        "to ensure immediate clinical, academic, and industrial readiness of our graduates:\n\n"
        "1. Academic Excellence & Transdisciplinary Courses: Trainees will complete two newly "
        "created graduate courses: 'ENV-701: Mathematical Environmental Modeling and Sensor Fusion' and "
        "'QML-720: Quantum-Inspired Machine Learning for Ecological Restorations'. These courses utilize "
        "the Hydrostar codebase for hands-on laboratory exercises.\n\n"
        "2. Mandatory Industrial & Agency Internships: Every MSc and PhD trainee will complete a "
        "4-month internship at partner organizations like Neuromorph Technologies Inc. or the Toronto "
        "and Region Conservation Authority (TRCA). This ensures trainees translate mathematical theory "
        "into field-ready environmental sensors and drone-based monitoring campaigns.\n\n"
        "3. Professional Skills & Regulatory Quality Systems: Trainees will undergo rigorous workshops "
        "in quality management systems (QMS), environmental data sovereignty, intellectual property, "
        "and technical communication. This includes learning to write structured Design History Files (DHF) "
        "and Device Master Records (DMR) for environmental sensors and hardware."
    )
    add_body(training_text)

    # ---------------------------------------------------------
    # BUDGET & FUNDING ALLOCATION
    # ---------------------------------------------------------
    add_heading_1("4. Budget and Funding Allocation Profile")
    budget_intro = (
        "In strict compliance with NSERC CREATE guidelines, over 80% of the requested $1,650,000 CAD "
        "funding is allocated directly to trainee stipends. Table 1 represents our 6-year financial profile."
    )
    add_body(budget_intro)

    # Create Budget Table
    # Rows: Category, Y1, Y2, Y3, Y4, Y5, Y6, Total
    budget_data = [
        ['Category', 'Year 1 ($)', 'Year 2 ($)', 'Year 3 ($)', 'Year 4 ($)', 'Year 5 ($)', 'Year 6 ($)', 'Total ($)'],
        ['MSc Stipends', '40,000', '60,000', '60,000', '60,000', '60,000', '60,000', '340,000'],
        ['PhD Stipends', '50,000', '120,000', '120,000', '120,000', '120,000', '120,000', '650,000'],
        ['PDF Stipends', '30,000', '60,000', '60,000', '60,000', '60,000', '60,000', '330,000'],
        ['Program Coord.', '15,000', '25,000', '25,000', '25,000', '25,000', '25,000', '140,000'],
        ['Travel / Field', '5,000', '15,000', '15,000', '15,000', '15,000', '15,000', '80,000'],
        ['Equipment / Mats', '8,000', '15,000', '15,000', '15,000', '15,000', '15,000', '83,000'],
        ['Workshops / Conf.', '2,000', '5,000', '5,000', '5,000', '5,000', '5,000', '27,000'],
        ['TOTAL', '150,000', '300,000', '300,000', '300,000', '300,000', '300,000', '1,650,000']
    ]

    table = doc.add_table(rows=len(budget_data), cols=len(budget_data[0]))
    add_table_borders(table)

    # Style cells
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = budget_data[r_idx][c_idx]
            
            # Formatting text size and alignment
            p = cell.paragraphs[0]
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            
            run = p.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            
            # Padding
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            
            # Header styling
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "0F172A") # #0f172a
            # Total row styling
            elif r_idx == len(budget_data) - 1:
                run.bold = True
                run.font.color.rgb = PRIMARY_COLOR
                set_cell_background(cell, "E2E8F0") # #e2e8f0
            else:
                set_cell_background(cell, "F8FAFC") # #f8fafc

    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(15)
    p_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = p_cap.add_run("Table 1: Proposed NSERC CREATE Hydrostar program budget allocation (in CAD).")
    run_cap.font.size = Pt(8.5)
    run_cap.italic = True
    run_cap.font.color.rgb = RGBColor(100, 116, 139)

    # ---------------------------------------------------------
    # EXPECTED IMPACT
    # ---------------------------------------------------------
    add_heading_1("5. Expected Impact and Trainee Outcomes")
    impact_text = (
        "Graduates from the Hydrostar CREATE program will play a crucial role in building Canada's "
        "clean-tech economy. By mastering transdisciplinary methodologies such as spatial Kalman sensor "
        "fusion, Gaussian species distribution modeling, continued fraction resource projections, and "
        "quantum approximate optimization, trainees will possess a unique set of skills that bridges "
        "advanced computer science and environmental restoration. Our industrial and agency partners "
        "have committed to first-priority interviews for all graduates, ensuring rapid translation of "
        "classroom research into tangible ecological impact for Canadian watersheds."
    )
    add_body(impact_text)

    # Save document
    doc.save(doc_path)
    print(f"✅ NSERC CREATE Grant Proposal Document successfully written to {doc_path}")

if __name__ == '__main__':
    generate_nserc_doc()
