
"""
Generates a Comprehensive Finite Math Technical Report for the Deep Brain Stimulation App.
Includes modules: FEA, Post Treatment Parameters, OCD, Depression, SAD, ASD, and Dementia.
Writes to .docx format.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import numpy as np
import io

def create_finite_math_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Comprehensive Finite Mathematical Framework\nfor Deep Brain Stimulation Systems', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Technical Report: Equations, Models, and Dynamics')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    doc.add_paragraph('Deep Brain Stimulation Scientific Advisory Board\nJanuary 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # Abstract
    doc.add_heading('Abstract', level=1)
    abstract_text = (
        "This technical report formalizes the mathematical underpinnings of the Deep Brain Stimulation (DBS) "
        "application. It covers the Finite Element Analysis (FEA) of electric fields, the coupled differential "
        "equations for OCD CSTC loops, the neurotransmitter dynamics in Depression, the quantum surface integrals "
        "and continued fraction metrics for ASD neural repair, and the neural field equations for Dementia pathology. "
        "Each module is grounded in stochastic physics and statistical mechanics."
    )
    doc.add_paragraph(abstract_text)

    # 1. FEA
    doc.add_heading('1. Finite Element Analysis (FEA) of Electric Fields and Post Treatment Parameters', level=1)
    doc.add_paragraph(
        "The fundamental physics of DBS involves the distribution of electric potential Phi in the brain tissue, "
        "governed by the quasi-static Poisson equation:"
    )
    doc.add_paragraph("∇ · (σ(x) ∇Φ(x)) = -∇ · J_source", style='Quote')
    doc.add_paragraph(
        "where sigma(x) represents the anisotropic conductivity tensor. The Volume of Tissue Activated (VTA) is the "
        "manifold where the second spatial derivative of the potential exceeds the axonal threshold:"
    )
    doc.add_paragraph("VTA = { x ∈ Ω | Δ²Φ(x) > θ_activation }", style='Quote')

    # 2. OCD
    doc.add_heading('2. OCD Module: CSTC Loop Dynamics', level=1)
    doc.add_paragraph(
        "OCD pathology is modeled as a hyperactive gain cycle in the Cortico-Striato-Thalamo-Cortical loop. "
        "The state vectors for Orbitofrontal Cortex (OFC), Caudate (C), and Thalamus (T) evolve according to coupled nonlinear ODEs:"
    )
    doc.add_paragraph("τ dC/dt = -C + S(w_OFC->C · OFC + I_DBS)", style='Quote')
    doc.add_paragraph("τ dT/dt = -T + S(w_C->T · C - w_GPi->T · GPi)", style='Quote')
    doc.add_paragraph("τ dOFC/dt = -OFC + S(w_T->OFC · T + I_ext)", style='Quote')
    
    # 3. Depression
    doc.add_heading('3. Depression Module: Neurotransmitter Dynamics', level=1)
    doc.add_paragraph(
        "The Depression model simulates the restoration of bioamines (Serotonin 5-HT, Dopamine DA) and the regulation of Glutamate (Glu). "
        "The dynamics under formulation are:"
    )
    doc.add_paragraph("dS_5HT/dt = α · E_eff(f) · (1 - S_5HT) - γ(S_5HT - S_base)", style='Quote')
    doc.add_paragraph(
        "The stimulation efficacy E_eff is strictly frequency-dependent, modeled as a Gaussian resonance around 130 Hz:"
    )
    doc.add_paragraph("E_eff(f) = A · exp( - (f - 130)² / 2σ² )", style='Quote')

    # Plot Depression
    fig, ax = plt.subplots(figsize=(6, 4))
    f = np.linspace(0, 200, 200)
    eff = np.exp(-((f - 130)**2) / (2 * 40**2))
    ax.plot(f, eff, 'b-', linewidth=2)
    ax.fill_between(f, eff, alpha=0.1, color='blue')
    ax.axvline(130, color='r', linestyle='--', label='130 Hz Optimum')
    ax.set_title("Depression Treatment Efficacy Profile")
    ax.set_xlabel("Frequency (Hz)")
    ax.set_ylabel("Normalized Efficacy")
    ax.legend()
    ax.grid(True, alpha=0.3)
    
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=150)
    plt.close()
    doc.add_picture(img_stream, width=Inches(5))

    # 4. SAD
    doc.add_heading('4. Seasonal Affective Disorder (SAD): Circadian Entrainment', level=1)
    doc.add_paragraph(
        "SAD is modeled as a desynchronization of the Circadian pacemaker (Suprachiasmatic Nucleus, SCN). "
        "DBS aims to re-entrain the phase φ(t) of the oscillator using a Kuramoto-like forcing term:"
    )
    doc.add_paragraph("dφ/dt = ω + K · sin(φ_dbs - φ) + Noise", style='Quote')
    doc.add_paragraph(
        "Where K is the coupling strength proportional to stimulation amplitude. "
        "Melatonin secretion M(t) is inversely coupled to SCN activity A_scn(t):"
    )
    doc.add_paragraph("dM/dt = -β · A_scn(t) · M + P_night(t)", style='Quote')

    # Plot SAD
    fig, ax = plt.subplots(figsize=(6, 4))
    t = np.linspace(0, 48, 200) # 48 hours
    scn = np.sin(0.26 * t) + 0.5 # SCN activity
    melatonin = np.maximum(0, np.cos(0.26 * t + np.pi)) # Anti-phase melatonin
    
    ax.plot(t, scn, 'r-', label='SCN Activity (Hz)')
    ax.plot(t, melatonin, 'b--', label='Melatonin (pg/mL)')
    ax.set_title("Circadian Rhythm Dynamics (SAD Model)")
    ax.set_xlabel("Time (Hours)")
    ax.set_ylabel("Normalized Amplitude")
    ax.legend()
    ax.grid(True, alpha=0.3)
    
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=150)
    plt.close()
    doc.add_picture(img_stream, width=Inches(5))


    # 5. ASD
    doc.add_heading('5. ASD Module: Quantum Neural Repair', level=1)
    doc.add_paragraph(
        "ASD repair focuses on restoring connectivity. We utilize a Quantum Surface Integral formalism where the "
        "neural connectivity is treated as a wavefunction Psi:"
    )
    doc.add_paragraph("Φ_repair = ∮_S Ψ* ∇Ψ · n dS", style='Quote')
    doc.add_paragraph(
        "The repair trajectory is modeled using Continued Fraction sequences, representing the hierarchical restructuring of neural pathways:"
    )
    doc.add_paragraph("C(t) = a0 + 1 / (a1 + 1 / (a2 + ...))", style='Quote')

    # 6. Dementia
    doc.add_heading('6. Dementia Module: Neural Decay & Memory', level=1)
    doc.add_paragraph(
        "Dementia pathology involves time-dependent atrophy and amyloid burden (Beta). The activity A of a memory region is governed by:"
    )
    doc.add_paragraph("τ dA/dt = -A + σ( Σ w_ij A_j + C_chol - P_beta )", style='Quote')
    
    # Save
    filename = "Comprehensive_DBS_Finite_Math_Report.docx"
    doc.save(filename)
    print(f"Report saved as {filename}")

if __name__ == "__main__":
    create_finite_math_report()
